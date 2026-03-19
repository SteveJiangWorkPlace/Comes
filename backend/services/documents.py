from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from pypdf import PdfReader


def extract_text_from_pdf(file_obj: BytesIO) -> str:
    reader = PdfReader(file_obj)
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def post_process_resume_data(data: dict[str, Any]) -> dict[str, Any]:
    forbidden_keywords = ["political", "cpc", "party member", "communist", "党员", "团员", "政治面貌"]

    def is_clean(text: Any) -> bool:
        return not any(k in str(text).lower() for k in forbidden_keywords)

    def fix_punct(text: str) -> str:
        out = text
        out = re.sub(r'["“]([^"”]+)[,]["”]', r"'\1',", out)
        out = re.sub(r'["“]([^"”]+)[\.]["”]', r"'\1'.", out)
        out = re.sub(r'["“]([^"”]+)[;]["”]', r"'\1';", out)
        if "Bachelor's" not in out and "Master's" not in out:
            out = re.sub(r'["“]([^"”]+)["”]', r"'\1'", out)
        return out

    def process_node(node: Any) -> Any:
        if isinstance(node, dict):
            return {k: process_node(v) for k, v in node.items()}
        if isinstance(node, list):
            return [process_node(i) for i in node]
        if isinstance(node, str):
            return fix_punct(node)
        return node

    if "education" in data:
        for edu in data["education"]:
            if "details" in edu:
                edu["details"] = [d for d in edu["details"] if is_clean(d)]

    if "additional_info" in data:
        data["additional_info"] = [d for d in data["additional_info"] if is_clean(d)]

    return process_node(data)


def sort_resume_data(data: dict[str, Any]) -> dict[str, Any]:
    def get_sort_key(item: dict[str, Any]) -> int:
        date_str = str(item.get("date", "")).lower()
        if any(x in date_str for x in ["present", "current", "now", "至今", "在职"]):
            return 9999
        years = re.findall(r"20\d{2}", date_str) or re.findall(r"19\d{2}", date_str)
        return int(years[-1]) if years else 0

    for sec in ["education", "professional_experience", "leadership_experience", "project_experience"]:
        if sec in data and isinstance(data[sec], list):
            data[sec].sort(key=get_sort_key, reverse=True)
    return data


def _set_table_borders_none(table) -> None:
    tbl_borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "nil")
        tbl_borders.append(border)
    table._tbl.tblPr.append(tbl_borders)


def _add_text(paragraph, text, bold=False, italic=False, underline=False, size=11) -> None:
    run = paragraph.add_run(str(text or ""))
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def _add_bottom_border(paragraph) -> None:
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    paragraph._p.get_or_add_pPr().append(p_bdr)


def _add_job_header(doc: Document, line1_left: str, line1_right: str, line2_left: str, line2_right: str) -> None:
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(5.5)
    table.columns[1].width = Inches(2.0)
    _set_table_borders_none(table)

    left_top = table.rows[0].cells[0].paragraphs[0]
    right_top = table.rows[0].cells[1].paragraphs[0]
    left_bottom = table.rows[1].cells[0].paragraphs[0]
    right_bottom = table.rows[1].cells[1].paragraphs[0]
    right_top.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_bottom.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_text(left_top, line1_left, bold=True)
    _add_text(right_top, line1_right)
    _add_text(left_bottom, line2_left, italic=True, underline=True)
    _add_text(right_bottom, line2_right)


def generate_resume_docx(data: dict[str, Any]) -> BytesIO:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    header = data.get("header", {})
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_text(p_name, header.get("name", "Name"), bold=True, size=20)

    contact_parts = [header.get("phone"), header.get("email"), header.get("address")]
    p_contact = doc.add_paragraph(" | ".join(x for x in contact_parts if x))
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(12)

    def add_section_header(title: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(title.upper())
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = True
        _add_bottom_border(p)

    if data.get("education"):
        add_section_header("Education")
        for edu in data["education"]:
            _add_job_header(doc, edu.get("school", ""), edu.get("location", ""), edu.get("degree", ""), edu.get("date", ""))
            for det in edu.get("details", []):
                p = doc.add_paragraph(style="List Bullet")
                if ":" in det:
                    left, right = det.split(":", 1)
                    _add_text(p, f"{left}:", bold=True)
                    _add_text(p, f" {right}")
                else:
                    _add_text(p, det)

    if data.get("professional_experience"):
        add_section_header("Professional Experience")
        for job in data["professional_experience"]:
            _add_job_header(doc, job.get("company", ""), job.get("location", ""), job.get("title", ""), job.get("date", ""))
            for bullet in job.get("bullets", []):
                p = doc.add_paragraph(style="List Bullet")
                _add_text(p, bullet)

    if data.get("project_experience"):
        add_section_header("Project Experience")
        for proj in data["project_experience"]:
            p = doc.add_paragraph()
            _add_text(p, proj.get("title", ""), bold=True)
            meta = doc.add_table(rows=1, cols=2)
            meta.columns[0].width = Inches(5.5)
            meta.columns[1].width = Inches(2.0)
            _set_table_borders_none(meta)
            _add_text(meta.cell(0, 0).paragraphs[0], proj.get("role", ""), italic=True, underline=True)
            right = meta.cell(0, 1).paragraphs[0]
            right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _add_text(right, proj.get("date", ""))
            if proj.get("research_topic"):
                p = doc.add_paragraph(style="List Bullet")
                _add_text(p, "Research Topic: ", bold=True)
                _add_text(p, proj.get("research_topic"))
            if proj.get("responsibilities"):
                p = doc.add_paragraph(style="List Bullet")
                _add_text(p, "Responsibilities: ", bold=True)
                _add_text(p, " ".join(str(x) for x in proj.get("responsibilities", [])))

    if data.get("leadership_experience"):
        add_section_header("Leadership & Activities")
        for act in data["leadership_experience"]:
            _add_job_header(
                doc,
                act.get("organization") or act.get("company", ""),
                act.get("location", ""),
                act.get("role") or act.get("title", ""),
                act.get("date", ""),
            )
            for bullet in act.get("bullets", []):
                p = doc.add_paragraph(style="List Bullet")
                _add_text(p, bullet)

    if data.get("honours"):
        add_section_header("Honours")
        for hon in data["honours"]:
            p = doc.add_paragraph(style="List Bullet")
            _add_text(p, hon)

    if data.get("additional_info"):
        add_section_header("Additional Information")
        for info in data["additional_info"]:
            p = doc.add_paragraph(style="List Bullet")
            if ":" in info:
                left, right = info.split(":", 1)
                _add_text(p, f"{left}:", bold=True)
                _add_text(p, f" {right}")
            else:
                _add_text(p, info)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

